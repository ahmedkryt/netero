from random import randint
from telegram.ext import Updater, CallbackQueryHandler, MessageHandler, Filters, CallbackContext
from telegram import InlineKeyboardMarkup, InlineKeyboardButton
from telegram.ext import CommandHandler
from threading import Thread
from docx import Document
import telegram
import datetime
import schedule
import logging
import time

# включаем логирование
logging.basicConfig(format='%(asctime)s - %(name)s - %(levelname)s - %(message)s', level=logging.INFO)

# для отправки сообщений с картинкой
def send_photo_with_caption(update, context, caption, markup=None, photo_path=None):
    if not photo_path:
        photo_path = 'https://drive.google.com/drive-viewer/AKGpihYy5z79VqmYw8FoXAbGjSgam-9G4xm8Zu80_h6JDbD01yB5UV7Xi5g2dsGwQsGbm5Xx3SxIb6w1D-QrLAzPehaxf8Fr9Sm1sA=s2560'
    chat_id = update.effective_chat.id
    context.bot.send_photo(chat_id=chat_id, photo=photo_path, caption=caption, reply_markup=markup)


# для отправки сообщений без картинки
def send_messages(update, context, text, markup=None):
    chat_id = update.effective_chat.id
    context.bot.send_message(chat_id=chat_id, text=text, reply_markup=markup)

# функция регулирующая задания, выдает на выходе рандомный идентификатор
def task(topic_name):
    within_topic = False
    identifiers = []
    next_is_identifier = False
    global decided_day_tasks
    global grade
    doc = Document('0. База заданий.docx')
    norm = 10 if grade == 10 else 20
    # если выполнено больше 10 заданий, не выдаем идентификатор
    try:
        if len(decided_day_tasks) >= norm:
            return 0
    except NameError:
        decided_day_tasks = []

    # Ищем раздел
    for paragraph in doc.paragraphs:
        if topic_name in paragraph.text:
            within_topic = True
        elif ("Раздел:" in paragraph.text) or ("Тема:" in paragraph.text):
            within_topic = False
        # Ищем Идентификатор
        if within_topic and ("Идентификатор:" in paragraph.text):
            next_is_identifier = True
        elif within_topic and next_is_identifier:
            # Собираем идентификаторы в список
            identifier = paragraph.text.strip()
            identifiers.append(identifier)
            next_is_identifier = False
    # Перебираем из доступного списка заданий те, что еще не решались за эти сутки
    task_number = identifiers[randint(0, len(identifiers) - 1)]
    while task_number in decided_day_tasks:
        task_number = identifiers[randint(0, len(identifiers))]
    # Записываем задание в список выполненных и отдаем его на выход
    decided_day_tasks.append(task_number)

    return task_number

# функция достает задание для тренажера из файла с заданиями
def get_task_text(topic_name, task_number, header):
    doc = Document('0. База заданий.docx')
    text = []
    topic_found = False
    task_found = False
    collecting_text = False
    for para in doc.paragraphs:
        # Проверяем название темы
        if para.text.strip() == topic_name:
            topic_found = True
            continue

        # Проверяем номер задания, если тема найдена
        if topic_found:
            if para.text.startswith(task_number):
                task_found = True
                continue

            if task_found:
                if para.text.startswith('Идентификатор'):
                    break  # Выходим из цикла, если нашли новое задание

                # Проверяем, является ли текущий абзац заголовком
                if para.text.strip() == header:
                    collecting_text = True  # Начинаем собирать текст
                    continue

                # Прерываем сбор текста при новом заголовке
                if any(para.text.startswith(h) for h in
                       ["Текст:", "Изображение:", "Ключ:", "Комментарий:"]) and collecting_text:
                    break  # Если нашли другой заголовок, выходим

                # Собираем текст
                if collecting_text:
                    if para.text.strip() != '':
                        text.append(para.text.strip())

    return '\n'.join(text)

# обнуляем список решенных заданий
def reset_tasks():
    global decided_day_tasks
    global decided_tasks
    for i in decided_day_tasks:
        decided_tasks.append(i)
    decided_day_tasks = []

# устанавливаем обнуление на 00:00 каждого дня
def decided_tasks_reset():
    schedule.every().day.at("00:00").do(reset_tasks)
    while True:
        schedule.run_pending()
        time.sleep(1)

# диалог установка таймера напоминаний
def timer(update, context, notice):
    keyboard = [
        [InlineKeyboardButton('Назад в меню', callback_data="В меню")],
        [InlineKeyboardButton('К задачам', callback_data="Попрактиковаться")]
    ]
    markup = InlineKeyboardMarkup(keyboard, one_time_keyboard=False)
    if notice:
        send_messages(update, context, '''Напиши удобное для тебя время в такой форме: __:__
Например, 09:00''')
    else:
        send_messages(update, context, 'Ну как хочешь 😔', markup)
        schedule.clear()


# сообщение для отправки напоминания
def send_reminder(context, chat_id: int):
    keyboard = [
        [InlineKeyboardButton('Согласиться', callback_data="Попрактиковаться")],
        [InlineKeyboardButton('Отказаться', callback_data="В меню")]
    ]
    markup = InlineKeyboardMarkup(keyboard, one_time_keyboard=False)
    context.bot.send_message(chat_id, text='Привет, пора позаниматься', reply_markup=markup)


# установка таймера
def set_reminder(update, context: CallbackContext, time):
    chat_id = update.effective_chat.id
    # удаляем предыдущую задачу, если она была
    schedule.clear()
    # добавляем новую задачу в планировщик
    schedule.every().day.at(time).do(send_reminder, context=context, chat_id=chat_id)


# цикл для таймера
def run_scheduler():
    while True:
        schedule.run_pending()
        time.sleep(1)



def handle_message(update, context):
    global user_status
    global task_key
    global decided_day_tasks
    # ввод пользователя для таймера
    if user_status == 'уведомления':
        keyboard = [
            [InlineKeyboardButton('Назад в меню', callback_data="В меню")],
            [InlineKeyboardButton('К задачам', callback_data="Попрактиковаться")]
        ]
        markup = InlineKeyboardMarkup(keyboard, one_time_keyboard=False)
        # время, которое ввел пользователь
        user_response = update.message.text
        # проверяем на валидность и отправляем на установку
        try:
            if datetime.datetime.strptime(user_response, '%H:%M').time():
                # если время записано без предшествующего нуля, корректируем данные
                if len(user_response.split(':')[0]) == 1:
                    user_response = '0' + user_response
                if len(user_response.split(':')[1]) == 1:
                    user_response = user_response[:-1] + '0' + user_response[-1]
                # устанавливаем
                set_reminder(update, context, user_response)
                # отвечаем пользователю
                context.bot.send_message(chat_id=update.effective_chat.id,
                                         text=f'Теперь точно не забудешь!', reply_markup=markup)
        except(IndexError, ValueError):
            update.message.reply_text(f'Не понимаю, попробуй еще раз', reply_markup=markup)
    # ввод пользователя в качестве ответа на задание
    elif user_status == 'тренажер':
        global grade
        norm = 10 if grade == 10 else 20
        # проверка ответа
        try:
            keyboard = [
                [InlineKeyboardButton('Следующее задание', callback_data="Попрактиковаться")]
            ]
            markup = InlineKeyboardMarkup(keyboard, one_time_keyboard=False)
            user_answer = update.message.text
            if '/' in task_key:
                key_values_list = task_key.split(" / ")
                print(key_values_list)
                print(user_answer)
                if user_answer.lower() in key_values_list:
                    text = (f'🟢 Верно!\n'
                            f'Ключ: {task_key}\n'
                            f'Выполнено {int(len(decided_day_tasks)/(norm/100))}% нормы')
                    context.bot.send_message(chat_id=update.effective_chat.id, text=text, reply_markup=markup)
                else:
                    text = (f'🔴 Неверно\n'
                            f'Ключ: {task_key}\n'
                            f'Выполнено {int(len(decided_day_tasks)/(norm/100))}% нормы')
                    context.bot.send_message(chat_id=update.effective_chat.id, text=text, reply_markup=markup)

            elif set(user_answer.lower()) == set(task_key.lower()):

                text = (f'🟢 Верно!\n'
                        f'Ключ: {task_key}\n'
                        f'Выполнено {int(len(decided_day_tasks)/(norm/100))}% нормы')
                context.bot.send_message(chat_id=update.effective_chat.id, text=text, reply_markup=markup)
            else:
                text = (f'🔴 Неверно\n'
                        f'Ключ: {task_key}\n'
                        f'Выполнено {int(len(decided_day_tasks)/(norm/100))}% нормы')
                context.bot.send_message(chat_id=update.effective_chat.id, text=text, reply_markup=markup)
        except(NameError):
            pass


# диалог Напоминания
def reminder(update, context):
    keyboard = [
        [InlineKeyboardButton('Настроить время уведомлений', callback_data="Время уведомлений")],
        [InlineKeyboardButton('Отказаться от уведомлений', callback_data="Отказаться от уведомлений")],
    ]
    markup = InlineKeyboardMarkup(keyboard, one_time_keyboard=False)
    send_photo_with_caption(update, context, 'Настроить уведомления', markup)


# Приветственное сообщение для первой сессии перед тренажером
def start_prac(update, context):
    global first_session
    first_session = False
    caption = '''График тренировок зависит от года обучения, так что признавайся!'''
    keyboard = [
        [InlineKeyboardButton('Я в 10 классе', callback_data="10 класс")],
        [InlineKeyboardButton('Я в 11 классе', callback_data="11 класс")]
        ]
    markup = InlineKeyboardMarkup(keyboard, one_time_keyboard=False)
    send_messages(update, context, caption, markup)


# Диалог Практика
def practice(update, context):
    keyboard = [
        [InlineKeyboardButton('Биология как наука', callback_data="Биология как наука практика")],
        [InlineKeyboardButton('Молекулярная и клеточная биология', callback_data="Молекулярная биология практика")],
        [InlineKeyboardButton('Генетика', callback_data="Генетика практика")],
        [InlineKeyboardButton('Биотехнология', callback_data="Биотехнология практика")],
        [InlineKeyboardButton('Селекция', callback_data="Селекция практика")],
        # [InlineKeyboardButton('Организм как биосистема', callback_data="Организм как биосистема практика")],
        # [InlineKeyboardButton('Многообразие органического мира', callback_data="Многообразие органического мира практика")],
        # [InlineKeyboardButton('Анатомия и физиология человека', callback_data="Анатомия и физиология человека практика")],
        # [InlineKeyboardButton('Эволюция', callback_data="Эволюция практика")],
        # [InlineKeyboardButton('Экология', callback_data="Экология практика")],
        [InlineKeyboardButton('В меню', callback_data="В меню")]
    ]
    markup = InlineKeyboardMarkup(keyboard, one_time_keyboard=False)
    send_messages(update, context, 'Выбери раздел', markup)

# Раздел Биология как наука Практика
def biology_as_science_prac(update, context):
    global task_key
    # получаем идентификатор случайного задания
    task_number = task('Биология как наука')
    # проверяем, выполнил ли пользователь норму на сутки
    if task_number == 0:
        send_messages(update, context, text='Норма выполнена!')
    else:
        task_key = get_task_text('Биология как наука', task_number, 'Ключ:')
        print(task_key)
        # проверяем, есть ли изображение в задании
        if get_task_text('Биология как наука', task_number, 'Изображение:'):
            send_photo_with_caption(update, context, get_task_text('Биология как наука', task_number, 'Текст:'), None,
                                    get_task_text('Биология как наука', task_number, 'Изображение:'))
        else:
            send_messages(update, context, text=get_task_text('Биология как наука', task_number, 'Текст:'))


# Раздел молекулярная и клеточная биология Практика
def molecular_and_cellular_biology_main_prac(update, context):
    keyboard = [
        [InlineKeyboardButton('История и методы цитологии. Клеточная теория', callback_data="История цитологии практика")],
        [InlineKeyboardButton('Биохимия', callback_data="Биохимия практика")],
        [InlineKeyboardButton('Строение клетки', callback_data="Строение клетки практика")],
        [InlineKeyboardButton('Метаболизм', callback_data="Метаболизм практика")],
        [InlineKeyboardButton('Матричные реакции', callback_data="Матричные реакции практика")],
        [InlineKeyboardButton('Клеточный цикл', callback_data="Клеточный цикл практика")]
    ]
    markup = InlineKeyboardMarkup(keyboard, one_time_keyboard=False)
    send_messages(update, context, text='Выбери тему', markup=markup)


# Раздел Генетика Практика
def genetics_prac(update, context):
    global task_key
    # получаем идентификатор случайного задания
    task_number = task('Генетика')
    # проверяем, выполнил ли пользователь норму на сутки
    if task_number == 0:
        send_messages(update, context, text='Хватит с тебя на сегодня. Беги отдыхать')
    else:
        task_key = get_task_text('Биология как наука', task_number, 'Ключ:')
        # проверяем, есть ли изображение в задании
        if get_task_text('Генетика', task_number, 'Изображение:'):
            send_photo_with_caption(update, context, get_task_text('Генетика', task_number, 'Текст:'), None,
                                    get_task_text('Генетика', task_number, 'Изображение:'))
        else:
            send_messages(update, context, text=get_task_text('Генетика', task_number, 'Текст:'))


# Раздел Биотехнология Практика
def biotechnology_prac(update, context):
    global task_key
    # получаем идентификатор случайного задания
    task_number = task('Биотехнология')
    # проверяем, выполнил ли пользователь норму на сутки
    if task_number == 0:
        send_messages(update, context, text='Хватит с тебя на сегодня. Беги отдыхать')
    else:
        task_key = get_task_text('Биология как наука', task_number, 'Ключ:')
        # проверяем, есть ли изображение в задании
        if get_task_text('Биотехнология', task_number, 'Изображение:'):
            send_photo_with_caption(update, context, get_task_text('Биотехнология', task_number, 'Текст:'), None,
                                    get_task_text('Биотехнология', task_number, 'Изображение:'))
        else:
            send_messages(update, context, text=get_task_text('Биотехнология', task_number, 'Текст:'))


# Раздел Селекция Практика
def breeding_prac(update, context):
    global task_key
    # получаем идентификатор случайного задания
    task_number = task('Селекция')
    # проверяем, выполнил ли пользователь норму на сутки
    if task_number == 0:
        send_messages(update, context, text='Хватит с тебя на сегодня. Беги отдыхать')
    else:
        task_key = get_task_text('Биология как наука', task_number, 'Ключ:')
        # проверяем, есть ли изображение в задании
        if get_task_text('Селекция', task_number, 'Изображение:'):
            send_photo_with_caption(update, context, get_task_text('Селекция', task_number, 'Текст:'), None,
                                    get_task_text('Селекция', task_number, 'Изображение:'))
        else:
            send_messages(update, context, text=get_task_text('Селекция', task_number, 'Текст:'))


# Тема История и методы цитологии Раздела Молекулярная и клеточная биология Практика
def history_and_methods_citology_prac(update, context):
    global task_key
    # получаем идентификатор случайного задания
    task_number = task('История и методы цитологии')
    # проверяем, выполнил ли пользователь норму на сутки
    if task_number == 0:
        send_messages(update, context, text='Хватит с тебя на сегодня. Беги отдыхать')
    else:
        task_key = get_task_text('Биология как наука', task_number, 'Ключ:')
        # проверяем, есть ли изображение в задании
        if get_task_text('История и методы цитологии', task_number, 'Изображение:'):
            send_photo_with_caption(update, context, get_task_text('История и методы цитологии', task_number, 'Текст:'), None,
                                    get_task_text('История и методы цитологии', task_number, 'Изображение:'))
        else:
            send_messages(update, context, text=get_task_text('История и методы цитологии', task_number, 'Текст:'))


# Тема Биохимия Раздела Молекулярная и клеточная биология Практика
def biochemistry_prac(update, context):
    global task_key
    # получаем идентификатор случайного задания
    task_number = task('Биохимия')
    # проверяем, выполнил ли пользователь норму на сутки
    if task_number == 0:
        send_messages(update, context, text='Хватит с тебя на сегодня. Беги отдыхать')
    else:
        task_key = get_task_text('Биология как наука', task_number, 'Ключ:')
        # проверяем, есть ли изображение в задании
        if get_task_text('Биохимия', task_number, 'Изображение:'):
            send_photo_with_caption(update, context, get_task_text('Биохимия', task_number, 'Текст:'), None,
                                    get_task_text('Биохимия', task_number, 'Изображение:'))
        else:
            send_messages(update, context, text=get_task_text('Биохимия', task_number, 'Текст:'))


# Тема Строение клетки Раздела Молекулярная и клеточная биология Практика
def cell_structure_prac(update, context):
    global task_key
    # получаем идентификатор случайного задания
    task_number = task('Строение клетки')
    # проверяем, выполнил ли пользователь норму на сутки
    if task_number == 0:
        send_messages(update, context, text='Хватит с тебя на сегодня. Беги отдыхать')
    else:
        task_key = get_task_text('Биология как наука', task_number, 'Ключ:')
        # проверяем, есть ли изображение в задании
        if get_task_text('Строение клетки', task_number, 'Изображение:'):
            send_photo_with_caption(update, context, get_task_text('Строение клетки', task_number, 'Текст:'), None,
                                    get_task_text('Строение клетки', task_number, 'Изображение:'))
        else:
            send_messages(update, context, text=get_task_text('Строение клетки', task_number, 'Текст:'))


# Тема Метаболизм Раздела Молекулярная и клеточная биология Практика
def metabolism_prac(update, context):
    global task_key
    # получаем идентификатор случайного задания
    task_number = task('Метаболизм')
    # проверяем, выполнил ли пользователь норму на сутки
    if task_number == 0:
        send_messages(update, context, text='Хватит с тебя на сегодня. Беги отдыхать')
    else:
        task_key = get_task_text('Биология как наука', task_number, 'Ключ:')
        # проверяем, есть ли изображение в задании
        if get_task_text('Метаболизм', task_number, 'Изображение:'):
            send_photo_with_caption(update, context, get_task_text('Метаболизм', task_number, 'Текст:'), None,
                                    get_task_text('Метаболизм', task_number, 'Изображение:'))
        else:
            send_messages(update, context, text=get_task_text('Метаболизм', task_number, 'Текст:'))


# Тема Матричные реакции Раздела Молекулярная и клеточная биология Практика
def matrix_reactions_prac(update, context):
    global task_key
    # получаем идентификатор случайного задания
    task_number = task('Матричные реакции')
    # проверяем, выполнил ли пользователь норму на сутки
    if task_number == 0:
        send_messages(update, context, text='Хватит с тебя на сегодня. Беги отдыхать')
    else:
        task_key = get_task_text('Биология как наука', task_number, 'Ключ:')
        # проверяем, есть ли изображение в задании
        if get_task_text('Матричные реакции', task_number, 'Изображение:'):
            send_photo_with_caption(update, context, get_task_text('Матричные реакции', task_number, 'Текст:'), None,
                                    get_task_text('Матричные реакции', task_number, 'Изображение:'))
        else:
            send_messages(update, context, text=get_task_text('Матричные реакции', task_number, 'Текст:'))


# Тема Клеточный цикл Раздела Молекулярная и клеточная биология Практика
def cell_cycle_prac(update, context):
    global task_key
    # получаем идентификатор случайного задания
    task_number = task('Клеточный цикл')
    # проверяем, выполнил ли пользователь норму на сутки
    if task_number == 0:
        send_messages(update, context, text='Хватит с тебя на сегодня. Беги отдыхать')
    else:
        task_key = get_task_text('Биология как наука', task_number, 'Ключ:')
        # проверяем, есть ли изображение в задании
        if get_task_text('Клеточный цикл', task_number, 'Изображение:'):
            send_photo_with_caption(update, context, get_task_text('Клеточный цикл', task_number, 'Текст:'), None,
                                    get_task_text('Клеточный цикл', task_number, 'Изображение:'))
        else:
            send_messages(update, context, text=get_task_text('Клеточный цикл', task_number, 'Текст:'))


# диалог Теория
def conspect(update, context):
    caption = '''❕ Советуем изучать конспекты по порядку — темы сверху помогут лучше разобраться в тех, что идут ниже.

Выбери раздел:'''
    keyboard = [
        [InlineKeyboardButton('Биология как наука', callback_data="Биология как наука теория")],
        [InlineKeyboardButton('Молекулярная и клеточная биология', callback_data="Молекулярная биология теория")],
        [InlineKeyboardButton('Генетика', callback_data="Генетика теория")],
        [InlineKeyboardButton('Биотехнология и селекция', callback_data="Биотехнология и Селекция теория")],
        [InlineKeyboardButton('Многообразие органического мира', callback_data="Многообразие мира теория")],
        [InlineKeyboardButton('Анатомия и физиология человека', callback_data="Анатомия и физиология теория")],
        [InlineKeyboardButton('Эволюция', callback_data="Эволюция теория")],
        [InlineKeyboardButton('Экология', callback_data="Экология теория")],
        [InlineKeyboardButton('В меню', callback_data="В меню")]
    ]
    markup = InlineKeyboardMarkup(keyboard, one_time_keyboard=False)
    send_messages(update, context, caption, markup)


# Раздел Биология как наука Теория
def biology_as_science_con(update, context):
    text = '''<a href='https://docs.google.com/document/d/1V8CRBD_NWHVGB5RVxbSXKdhwsmAe_irgcWOD-T3M7_Q/edit?usp=drive_link'>📑 Свойства и уровни организации биосистем. Разделы и методы биологии</a>

<a href='https://docs.google.com/document/d/17gU7Co2qcBKVhptiXlLvfKTOoh_zsnWFx-De6lYIdiQ/edit?usp=drive_link'>📑 Исследование</a>'''
    chat_id = update.effective_chat.id
    context.bot.send_message(chat_id=chat_id, text=text, parse_mode=telegram.ParseMode.HTML, disable_web_page_preview = True)


# Раздел молекулярная и клеточная биология Теория
def molecular_and_cellular_biology_main_con(update, context):
    keyboard = [
        [InlineKeyboardButton('История и методы цитологии. Клеточная теория', callback_data="История цитологии теория")],
        [InlineKeyboardButton('Биохимия', callback_data="Биохимия теория")],
        [InlineKeyboardButton('Строение клетки', callback_data="Строение клетки теория")],
        [InlineKeyboardButton('Метаболизм', callback_data="Метаболизм теория")],
        [InlineKeyboardButton('Матричные реакции', callback_data="Матричные реакции теория")],
        [InlineKeyboardButton('Клеточный цикл', callback_data="Клеточный цикл теория")]
    ]
    markup = InlineKeyboardMarkup(keyboard, one_time_keyboard=False)
    send_messages(update, context, text='Выбери тему', markup=markup)


# Раздел Генетика Теория
def genetics_con(update, context):
    text = '''<a href='https://docs.google.com/document/d/15lipZln844hOfzfO5ZQfXJwZ_E1cjO3ciOAKEukInj0/edit?usp=sharing'>📑 Основные генетические понятия и символы. Методы генетики</a>

<a href='https://docs.google.com/document/d/1I6EyUTvBh4OW5s_Kfw4Cflx8gQmi9J_-DgDInYyf9vo/edit?usp=sharing'>📑 Генетика Менделя</a>

<a href = 'https://docs.google.com/document/d/1C4Wk499yuDfAhM4welcv_TCRn_-Q255GwP7CRCnnCQw/edit?usp=sharing'>📑 Генетика Моргана. Хромосомная теория наследственности</a>

<a href = 'https://docs.google.com/document/d/19FIgIge7H-xlA1a8S-98keyqapXTaNXin-B0MZgdCK0/edit?usp=sharing'>📑 Изменчивость. Генетические болезни</a>'''
    chat_id = update.effective_chat.id
    context.bot.send_message(chat_id=chat_id, text=text, parse_mode=telegram.ParseMode.HTML,
                             disable_web_page_preview=True)


# Раздел Биотехнология и Селекция Теория
def biotechnology_con(update, context):
    text = '''<a href='https://docs.google.com/document/d/1ipgZOoMw1KO1PKInvnStYFLgU7xI2F5ASkWpXyx1Tdo/edit?usp=sharing'>📑 Биотехнология</a>
<a href='https://docs.google.com/document/d/16cYkwZXNVub2ylcaCYR5E1m4F8Pdx8BAQ7YR7-1KTlA/edit?usp=drive_link'>📑 Селекция</a>'''
    chat_id = update.effective_chat.id
    context.bot.send_message(chat_id=chat_id, text=text, parse_mode=telegram.ParseMode.HTML,
                             disable_web_page_preview=True)


# Тема История и методы цитологии Раздела Молекулярная и клеточная биология Теория
def history_and_methods_citology_con(update, context):
    text = '''<a href='https://docs.google.com/document/d/1kA8Ue_5iLpen2GmOEIAfJByZfJxufh6R_xhY5Wavk10/edit?usp=sharing'>📑 История и методы цитологии. Клеточная теория</a>'''
    chat_id = update.effective_chat.id
    context.bot.send_message(chat_id=chat_id, text=text, parse_mode=telegram.ParseMode.HTML,
                             disable_web_page_preview=True)


# Тема Биохимия Раздела Молекулярная и клеточная биология Теория
def biochemistry_con(update, context):
    text = '''<a href='https://docs.google.com/document/d/1v1nQ7noVqfkrt0gAp0QU8-9aafwNaK4zZbj2Y2OW5a4/edit?usp=sharing'>📑 Химический состав клетки. Вода. Минеральные вещества</a>

<a href='https://docs.google.com/document/d/1AuySgB2dddxSIPCRsaFn2S014y8v1jSBP1KJJ_iAvsA/edit?usp=sharing'>📑 Липиды</a>

<a href = 'https://docs.google.com/document/d/143tRQQUbF-QWQdfDAvV3ny2GL6JLM3yD6fQaaR3b3Ks/edit?usp=sharing'>📑 Углеводы</a>

<a href = 'https://docs.google.com/document/d/16WXZRksgYRACQrLb70Ok_4wXdOtfqlaKWXGAWnMqJUk/edit?usp=sharing'>📑 Нуклеиновые кислоты. АТФ</a>

<a href = 'https://docs.google.com/document/d/1CMBsKh9_PdMzpeK5WvObAXNKlDtqqh3mYM6_nQA6ul0/edit?usp=sharing'>📑 Белки. Ферменты</a>'''
    chat_id = update.effective_chat.id
    context.bot.send_message(chat_id=chat_id, text=text, parse_mode=telegram.ParseMode.HTML,
                             disable_web_page_preview=True)


# Тема Строение клетки Раздела Молекулярная и клеточная биология Теория
def cell_structure_con(update, context):
    text = '''<a href='https://docs.google.com/document/d/1NGdosXt8mkpFgRpK7Bf0Ry5XVdGeHiH1-6WWTBSgh_I/edit?usp=sharing'>📑 Прокариоты и эукариоты</a>

<a href='https://docs.google.com/document/d/1nON2sdHiu0wL7JLZtvbXkRmwb53eqMAavBYgAK3YjoI/edit?usp=sharing'>📑 Строение прокариотической клетки. Бактерии и археи</a>

<a href = 'https://docs.google.com/document/d/15x990_dywHVmcpIQncvdmKLLCV8saOZFSNHLBZGOUeA/edit?usp=sharing'>📑 Строение эукариотической клетки</a>'''

    chat_id = update.effective_chat.id
    context.bot.send_message(chat_id=chat_id, text=text, parse_mode=telegram.ParseMode.HTML,
                             disable_web_page_preview=True)


# Тема Метаболизм Раздела Молекулярная и клеточная биология Теория
def metabolism_con(update, context):
    text = '''<a href='https://docs.google.com/document/d/1JHqaKDWGyHykpTrSjzcrBHqUV4ClDk031kj0CLTvtqU/edit?usp=sharing'>📑 Метаболизм</a>

<a href='https://docs.google.com/document/d/14wt_qHqAuedkzwGs2znmR43QKeUKQIVH7aNPyLhG1-M/edit?usp=sharing'>📑 Анаболизм: хемосинтез и фотосинтез</a>

<a href='https://docs.google.com/document/d/1bZeJRZv4es3vVe7DQuVOJUY2FcAuxVLz7Ob1ruNXIdg/edit?usp=sharing'>📑 Катаболизм: брожение и клеточное дыхание</a>'''
    chat_id = update.effective_chat.id
    context.bot.send_message(chat_id=chat_id, text=text, parse_mode=telegram.ParseMode.HTML,
                             disable_web_page_preview=True)


# Тема Матричные реакции Раздела Молекулярная и клеточная биология Теория
def matrix_reactions_con(update, context):
    text = '''<a href='https://docs.google.com/document/d/1_O0of-Q1Ngw-BuosM-SVtHt09Cp7EyzMhyeWA07dOTc/edit?usp=sharing'>📑 Репликация ДНК</a>

<a href='https://docs.google.com/document/d/19ki2hZ7bfwCil6pzt58jxPuSPQKzqKxj1EtZ-GKNOvw/edit?usp=sharing'>📑 Транскрипция и трансляция</a>'''
    chat_id = update.effective_chat.id
    context.bot.send_message(chat_id=chat_id, text=text, parse_mode=telegram.ParseMode.HTML,
                             disable_web_page_preview=True)


# Тема Клеточный цикл Раздела Молекулярная и клеточная биология Теория
def cell_cycle_con(update, context):
    text = '''<a href='https://docs.google.com/document/d/18DfCOdE0V7TmNSV3pYbMLxmkFgQknoLauC6du38_5HM/edit?usp=sharing'>📑 Клеточный цикл</a>'''
    chat_id = update.effective_chat.id
    context.bot.send_message(chat_id=chat_id, text=text, parse_mode=telegram.ParseMode.HTML,
                             disable_web_page_preview=True)


# обработчик кнопок
def button(update, context):
    query = update.callback_query
    variant = query.data
    query.answer()
    global user_status
    global first_session
    global grade
    if variant == 'Преамбула':
        advent(update, context)
    if variant == '10 класс':
        send_photo_with_caption(update, context, '''🐣 Твоя норма — 10 заданий каждый день. 
Чтобы не забывать о подготовке, можешь настроить уведомления. Я или Ив будем тебе писать.''')
        grade = 10
        time.sleep(2)
        reminder(update, context)
    if variant == '11 класс':
        send_photo_with_caption(update, context, '''🐥 Твоя норма — 20 заданий каждый день.
Чтобы не забывать о подготовке, можешь настроить уведомления. Я или Ив будем тебе писать.''')
        grade = 11
        time.sleep(2)
        reminder(update, context)

    if variant == 'Разобраться с теорией':
        user_status = 'теория'
        conspect(update, context)
    if variant == 'Попрактиковаться':
        user_status = 'тренажер'
        if first_session:
            start_prac(update, context)
        else:
            practice(update, context)
    if variant == 'Настроить уведомления':
        reminder(update, context)

    if variant == 'В меню':
        user_status = 'menu'
        menu(update, context)

    if variant == 'Время уведомлений':
        user_status = 'уведомления'
        timer(update, context, True)
    if variant == 'Отказаться от уведомлений':
        timer(update, context, False)
    # теория
    if variant == 'Биология как наука теория':
        biology_as_science_con(update, context)
    if variant == 'Молекулярная биология теория':
        molecular_and_cellular_biology_main_con(update, context)
    if variant == 'История цитологии теория':
        history_and_methods_citology_con(update, context)
    if variant == 'Биохимия теория':
        biochemistry_con(update, context)
    if variant == 'Строение клетки теория':
        cell_structure_con(update, context)
    if variant == 'Метаболизм теория':
        metabolism_con(update, context)
    if variant == 'Матричные реакции теория':
        matrix_reactions_con(update, context)
    if variant == 'Клеточный цикл теория':
        cell_cycle_con(update, context)
    if variant == 'Генетика теория':
        genetics_con(update, context)
    if variant == 'Биотехнология и Селекция теория':
        biotechnology_con(update, context)
    # практика
    if variant == 'Биология как наука практика':
        biology_as_science_prac(update, context)
    if variant == 'Молекулярная биология практика':
        molecular_and_cellular_biology_main_prac(update, context)
    if variant == 'История цитологии практика':
        history_and_methods_citology_prac(update, context)
    if variant == 'Биохимия практика':
        biochemistry_prac(update, context)
    if variant == 'Строение клетки практика':
        cell_structure_prac(update, context)
    if variant == 'Метаболизм практика':
        metabolism_prac(update, context)
    if variant == 'Матричные реакции практика':
        matrix_reactions_prac(update, context)
    if variant == 'Клеточный цикл практика':
        cell_cycle_prac(update, context)
    if variant == 'Генетика практика':
        genetics_prac(update, context)
    if variant == 'Биотехнология практика':
        biotechnology_prac(update, context)
    if variant == 'Селекция практика':
        breeding_prac(update, context)


# основное окно
def menu(update, context):
    caption = 'Чем займемся?'
    keyboard = [
        [InlineKeyboardButton('Изучить конспект', callback_data="Разобраться с теорией")],
        [InlineKeyboardButton('Потренировать знание теории', callback_data="Zero")],
        [InlineKeyboardButton('Потренировать решение заданий', callback_data="Попрактиковаться")],
        [InlineKeyboardButton('Пройти квест', callback_data="Zero")],
        [InlineKeyboardButton('Настроить уведомления', callback_data="Настроить уведомления")],
        [InlineKeyboardButton('Почитать комикс о важном', callback_data="Zero")],
        [InlineKeyboardButton('Оставить фидбэк', callback_data="Zero")]
    ]
    markup = InlineKeyboardMarkup(keyboard, one_time_keyboard=False)
    send_photo_with_caption(update, context, caption, markup)


# Точка входа
def start(update, context):
    global first_session
    keyboard = [
            [InlineKeyboardButton('Погодите, как вы собрались меня готовить?', callback_data="Преамбула")]
    ]
    markup = InlineKeyboardMarkup(keyboard, one_time_keyboard=False)
    name = update.message.chat.first_name
    first_session = True
    caption = '''Привет, {name}! 👋
    
Если ты ищешь лучший способ подготовиться к ЕГЭ по биологии — поздравляем, вот он!'''.format(name=name)
    send_photo_with_caption(update, context, caption, markup)


def advent(update, context):
    caption = '''Для подготовки у нас есть:
    
🍀 Конспекты по всем темам кодификатора
🍀 Тренажёр для отработки заданий из банка ФИПИ
🍀 Квесты вместо тестов'''
    send_photo_with_caption(update, context, caption)
    caption = '''Скажем сразу: ЕГЭ по биологии — серьёзный экзамен. К нему не подготовиться с нуля на 90+ баллов за месяц.

На освоение теории у тебя уйдёт ~Х часов, на отработку заданий — ~Y часов. Ещё нужно заложить пару месяцев на восполнение пробелов в знаниях и прорешивание вариантов из сборника ФИПИ.

🍀 Если ты в 10 классе, изучай по одному конспекту в неделю и каждый день решай 10 заданий по теме конспекта (и всем изученным прежде)
🍀 Если ты в 11 классе, изучай по два конспекта в неделю и каждый день решай 20 заданий по темам конспектов (и всем изученным прежде)'''
    send_photo_with_caption(update, context, caption)
    keyboard = [
        [InlineKeyboardButton('Было бы славно (Начать квест)', callback_data="Квест Первый день")],
        [InlineKeyboardButton('В меню', callback_data='В меню')]
    ]
    markup = InlineKeyboardMarkup(keyboard, one_time_keyboard=False)
    caption = '''В общем-то, это все инструкции. Приступай хоть сейчас.

А впрочем, у нас есть ещё 20 минут до начала рабочего дня, можем повести тебе экскурсию по БиБу и заодно познакомиться поближе. Что скажешь?'''
    send_photo_with_caption(update, context, caption, markup)


# основная
def main():
    updater = Updater("7680944125:AAFOtGfVXlE8rN3yrbxS6vwRYuddmEJEX7Q", use_context=True)
    dp = updater.dispatcher
    scheduler_thread = Thread(target=run_scheduler, daemon=True)
    scheduler_thread.start()
    thread = Thread(target=decided_tasks_reset)
    thread.start()
    dp.add_handler(CommandHandler("start", start))
    dp.add_handler(CallbackQueryHandler(button))
    dp.add_handler(MessageHandler(Filters.text, handle_message))
    updater.start_polling()
    updater.idle()


if __name__ == '__main__':
    main()
